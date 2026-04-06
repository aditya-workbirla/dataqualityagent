import streamlit as st
import os
import io
import pandas as pd
import os
from dotenv import load_dotenv # Ensure the local modules can be found
from agents.agent import build_data_quality_graph, AgentState
from ui import ui_components
from agents import kg_builder
import streamlit.components.v1 as components
from langgraph.checkpoint.sqlite import SqliteSaver
import sqlite3
import uuid
import json

# Load environment variables
load_dotenv(override=True)

st.set_page_config(
    page_title="DataQA — Agentic Data Quality Platform",
    page_icon="🔍",
    layout="wide"
)

# Load custom CSS
try:
    with open("ui/style.css", "r") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
except FileNotFoundError:
    st.warning("Could not load ui/style.css. Ensure the file exists.")

# Top Navigation Bar mimicking HTML
st.markdown(ui_components.get_nav_html(), unsafe_allow_html=True)

# Ensure API keys are set up
use_azure = os.getenv("USE_AZURE_OPENAI", "false").lower() == "true"
api_key_set = False

if use_azure:
    if "AZURE_OPENAI_API_KEY" in os.environ:
        api_key_set = True
    else:
        st.error("Warning: `AZURE_OPENAI_API_KEY` is not set in your environment.")
else:
    if "OPENAI_API_KEY" in os.environ:
        api_key_set = True
    else:
        st.error("Warning: `OPENAI_API_KEY` is not set in your environment.")

if not api_key_set:
    st.info("Please set your API key in the `.env` file and restart the application.")
    st.stop()

# ── LLM endpoint connectivity pre-check ──────────────────────────────────────
def _check_llm_endpoint() -> tuple[bool, str, float]:
    """
    Quick TCP-level check (≤5 s) + a live ping to the completions endpoint.
    Returns (ok, message, latency_ms).
    """
    import socket, time, httpx
    if os.getenv("USE_AZURE_OPENAI", "false").lower() == "true":
        raw = os.getenv("AZURE_OPENAI_ENDPOINT", "")
    else:
        raw = "https://api.openai.com"

    from urllib.parse import urlparse
    parsed = urlparse(raw)
    host = parsed.hostname or raw
    port = parsed.port or (443 if parsed.scheme == "https" else 80)

    # Step 1: TCP reachability
    t0 = time.time()
    try:
        sock = socket.create_connection((host, port), timeout=5)
        sock.close()
        tcp_ms = (time.time() - t0) * 1000
    except Exception as e:
        return False, f"TCP connect to `{raw}` failed — {type(e).__name__}: {e}", 0.0

    # Step 2: live API ping (1 token)  
    _is_private = any(host.startswith(p) for p in ("10.", "172.", "192.168."))
    host_hdr = os.getenv("AZURE_OPENAI_HOST_HEADER", "")
    api_key   = os.getenv("AZURE_OPENAI_API_KEY", "") or os.getenv("OPENAI_API_KEY", "")
    deploy    = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")
    api_ver   = os.getenv("AZURE_OPENAI_API_VERSION", "2023-05-15")

    if os.getenv("USE_AZURE_OPENAI", "false").lower() == "true":
        url = f"{raw.rstrip('/')}/openai/deployments/{deploy}/chat/completions?api-version={api_ver}"
        headers = {"api-key": api_key, "Content-Type": "application/json"}
        if _is_private and host_hdr:
            headers["Host"] = host_hdr
    else:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    t1 = time.time()
    try:
        r = httpx.post(
            url,
            json={"messages": [{"role": "user", "content": "Say OK"}], "max_completion_tokens": 10},
            headers=headers,
            timeout=httpx.Timeout(12.0, connect=5.0),
            verify=False,
            trust_env=not _is_private,
        )
        api_ms = (time.time() - t1) * 1000
        if r.status_code == 200:
            return True, f"Connected · `{raw}` · TCP {tcp_ms:.0f} ms · API {api_ms:.0f} ms", api_ms
        else:
            err = r.json().get("error", {}).get("message", r.text[:120])
            return False, f"HTTP {r.status_code} from `{raw}` — {err}", api_ms
    except Exception as e:
        return False, f"TCP OK but API call failed — {type(e).__name__}: {e}", 0.0

if "llm_endpoint_ok" not in st.session_state:
    with st.sidebar:
        with st.status("🔌 Checking LLM endpoint…", expanded=False) as _ep_status:
            ok, reason, latency = _check_llm_endpoint()
            st.session_state["llm_endpoint_ok"] = ok
            st.session_state["llm_endpoint_reason"] = reason
            if ok:
                _ep_status.update(label=f"✅ LLM endpoint reachable", state="complete", expanded=False)
            else:
                _ep_status.update(label="❌ LLM endpoint unreachable", state="error", expanded=True)
                st.error(reason)

# Sidebar persistent endpoint badge
with st.sidebar:
    if st.session_state.get("llm_endpoint_ok"):
        st.success(f"🟢 **LLM Online** — {st.session_state['llm_endpoint_reason']}", icon=None)
    else:
        st.error(
            f"🔴 **LLM Offline**\n\n{st.session_state.get('llm_endpoint_reason','')}\n\n"
            f"**Possible fixes:**\n"
            f"- Check `AZURE_OPENAI_ENDPOINT` and `AZURE_OPENAI_API_KEY` in `.env`\n"
            f"- Verify `AZURE_OPENAI_DEPLOYMENT_NAME` matches the deployed model name in Azure\n"
            f"- Check `AZURE_OPENAI_API_VERSION` is supported by your deployment\n"
            f"- If endpoint is a private IP (`10.x.x.x`), ensure subnet `10.10.11.x` is whitelisted in Azure Networking"
        )
    if st.button("🔄 Re-check endpoint", use_container_width=True):
        # Clear cached result so the check runs again on the next rerun
        for k in ("llm_endpoint_ok", "llm_endpoint_reason"):
            st.session_state.pop(k, None)
        st.rerun()
    st.divider()

# Initialize session state variables if they don't exist
if "report" not in st.session_state:
    st.session_state["report"] = None
if "annotated_df" not in st.session_state:
    st.session_state["annotated_df"] = None
if "analysis_done" not in st.session_state:
    st.session_state["analysis_done"] = False
if "agent_state" not in st.session_state:
    st.session_state["agent_state"] = None
if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []
if "chat_execution_logs" not in st.session_state:
    # Keyed by chat turn index (int) → dict with plan + function execution details
    st.session_state["chat_execution_logs"] = {}
if "checkpointer" not in st.session_state:
    conn = sqlite3.connect("database/app.db", check_same_thread=False)
    saver = SqliteSaver(conn)
    saver.setup()  # Ensures the checkpoints tables are created
    st.session_state["checkpointer"] = saver
if "thread_id" not in st.session_state:
    st.session_state["thread_id"] = str(uuid.uuid4())

def convert_df_to_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")

def convert_report_to_docx(report_md: str) -> bytes:
    """
    Converts a Markdown report string to a .docx file in memory.
    Headings (#, ##, ###) become Word headings; **bold** becomes bold runs;
    everything else becomes normal paragraphs.
    """
    from docx import Document
    from docx.shared import Pt
    import re

    doc = Document()

    # Style the default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_run_with_bold(para, text: str):
        """Split text on **bold** markers and add runs accordingly."""
        parts = re.split(r"\*\*(.*?)\*\*", text)
        for i, part in enumerate(parts):
            run = para.add_run(part)
            if i % 2 == 1:   # odd index = inside ** **
                run.bold = True

    for line in report_md.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_bold(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_run_with_bold(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        else:
            p = doc.add_paragraph()
            add_run_with_bold(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# Sidebar Session History
with st.sidebar:
    st.header("🗄️ Session History")
    # Fetch threads from sqlite checkpointer
    try:
        with sqlite3.connect("database/app.db", timeout=5.0) as cp_conn:
            cp_cursor = cp_conn.cursor()
            cp_cursor.execute("SELECT DISTINCT thread_id FROM checkpoints")
            threads = [row[0] for row in cp_cursor.fetchall() if row[0] != "main_thread"]
    except Exception:
        threads = []
        
    if threads:
        thread_names = {"-- New Session --": "-- New Session --"}
        try:
            app_for_history = build_data_quality_graph(checkpointer=st.session_state["checkpointer"])
            for t in threads:
                config = {"configurable": {"thread_id": t}}
                state = app_for_history.get_state(config).values
                if state and "user_context_prompt" in state:
                    prompt = state["user_context_prompt"].strip().replace('\n', ' ')
                    name = (prompt[:35] + "...") if len(prompt) > 35 else prompt
                    if not name:
                         name = "Empty Prompt"
                    # Add short UUID tag to differentiate identical prompts
                    thread_names[t] = f"{name} ({t[:4]})"
                else:
                    thread_names[t] = t
        except Exception:
            for t in threads: thread_names[t] = t
            
        selected_thread = st.selectbox(
            "Resume Previous Session", 
            options=["-- New Session --"] + threads,
            format_func=lambda x: thread_names.get(x, x)
        )
        if selected_thread != "-- New Session --" and selected_thread != st.session_state.get("thread_id"):
            st.session_state["thread_id"] = selected_thread
            config = {"configurable": {"thread_id": selected_thread}}
            app = build_data_quality_graph(checkpointer=st.session_state["checkpointer"])
            saved_state = app.get_state(config).values
            
            if saved_state:
                st.session_state["report"] = saved_state.get("report", "No report generated.")
                try:
                    df_json_str = saved_state.get("df_json", "[]")
                    st.session_state["annotated_df"] = pd.read_json(io.StringIO(df_json_str))
                except Exception:
                    st.session_state["annotated_df"] = pd.DataFrame()
                st.session_state["analysis_done"] = True
                st.session_state["agent_state"] = saved_state
                
                chat_hist = []
                if "messages" in saved_state:
                    from langchain_core.messages import HumanMessage, AIMessage
                    for m in saved_state["messages"]:
                        if isinstance(m, HumanMessage):
                            chat_hist.append({"role": "user", "content": m.content})
                        elif isinstance(m, AIMessage) and m.content:
                            chat_hist.append({"role": "assistant", "content": m.content})
                st.session_state["chat_history"] = chat_hist
                st.session_state["chat_execution_logs"] = {}  # logs don't persist across sessions
                
                # Rebuild Knowledge Graph since it's not natively in state
                if st.session_state["annotated_df"] is not None:
                    try:
                        st.session_state["kg_html"] = kg_builder.build_knowledge_graph(st.session_state["annotated_df"], "data/knowledge_graph.json")
                    except Exception:
                        pass
                
                st.rerun()

tab1, tab2, tab3, tab4, tab5 = st.tabs(["New Analysis", "Functions DB", "Knowledge Graph", "Domain Knowledge Base", "Agent Workflow"])

with tab1:
    st.markdown(ui_components.get_header_html(), unsafe_allow_html=True)
    
    st.markdown('<div class="card"><div class="card-label">Dataset Upload</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Drop dataset here or click to browse", type=["csv", "xlsx", "xls", "parquet"], label_visibility="collapsed")
    st.markdown('</div><div class="spacer-20"></div>', unsafe_allow_html=True)
    
    df = None
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)

            # Tell execute_writer the real filename so execute.py loads the right file
            try:
                from agents.execute_writer import set_dataset_path
                set_dataset_path(uploaded_file.name)
            except Exception:
                pass

            st.markdown(f"""
            <div class="file-preview visible" style="margin-top:-10px; margin-bottom: 20px;">
              <div class="file-icon">📊</div>
              <div class="file-info">
                <div class="file-name">{uploaded_file.name}</div>
                <div class="file-size">{len(df):,} rows · {len(df.columns)} columns</div>
              </div>
              <div class="badge ok">Ready</div>
            </div>
            """, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Error loading file: {e}")
            st.stop()
    
    st.markdown('<div class="card"><div class="card-label">Dataset Context + Requirements</div>', unsafe_allow_html=True)
    user_context = st.text_area(
        "Prompt", 
        value="General dataset.",
        placeholder="Describe your dataset and what it represents...\ne.g. 'This is a production log from a Kraft pulp and fiber processing plant...'",
        label_visibility="collapsed",
        height=120
    )
    st.markdown('</div><div class="spacer-20"></div>', unsafe_allow_html=True)
    
    run_button = st.button("▶ Run Quality Analysis", type="primary", use_container_width=False)

    # ── Node → human-readable label map ──────────────────────────────────────
    _NODE_LABELS = {
        "collect_function_results": ("⚙️", "Running pre-defined quality checks on dataset…"),
        "check_existing_kb":        ("🗄️", "Checking for existing domain knowledge base…"),
        "knowledge_agent":          ("🔍", "Building domain knowledge base (Process / Physics / Equipment / OEM)…"),
        "critique_agent":           ("🧐", "Critiquing & scoring the knowledge base…"),
        "finalize_kb":              ("💾", "Finalising & saving the knowledge base…"),
        "chat_planner":             ("📋", "Planner agent is mapping out execution steps…"),
        "rag_retrieval":            ("🔎", "Retrieving relevant knowledge base sections (RAG)…"),
        "quality_analyst":          ("🧠", "Reasoning agent is analysing data quality…"),
        "tool_execution":           ("🔧", "Executing custom data quality function…"),
        "generate_report":          ("📋", "Generating the final quality report…"),
    }

    if run_button and df is not None:
        st.session_state["thread_id"] = str(uuid.uuid4())
        st.session_state["chat_history"] = []
        st.session_state["chat_execution_logs"] = {}

        app_graph = build_data_quality_graph(checkpointer=st.session_state["checkpointer"])
        config    = {"configurable": {"thread_id": st.session_state["thread_id"]}}
        initial_state = {
            "df_json": df.to_json(orient="records"),
            "user_context_prompt": user_context,
            "function_results_summary": {}, "messages": [],
            "issues": [], "bad_indices_per_column": {}, "report": "",
        }

        final_state = None
        error_msg   = None

        with st.status("🚀 Starting analysis pipeline…", expanded=True) as run_status:
            step_placeholder = st.empty()
            try:
                tool_call_count = 0
                for chunk in app_graph.stream(initial_state, config, stream_mode="updates"):
                    # chunk is {node_name: state_delta}
                    for node_name, node_output in chunk.items():
                        icon, label = _NODE_LABELS.get(node_name, ("🔄", f"Running `{node_name}`…"))

                        # Special-case: tool_execution — show which tool was called
                        if node_name == "tool_execution":
                            tool_call_count += 1
                            msgs = node_output.get("messages", [])
                            tool_names = [
                                m.get("name", "tool") if isinstance(m, dict) else getattr(m, "name", "tool")
                                for m in msgs
                            ]
                            tool_str = ", ".join(set(tool_names)) or "custom function"
                            # Check if Code Tester auto-corrected anything
                            auto_fixes = []
                            for m in msgs:
                                raw_c = m.get("content", "") if isinstance(m, dict) else getattr(m, "content", "")
                                try:
                                    _r = json.loads(raw_c) if isinstance(raw_c, str) else {}
                                    if _r.get("code_auto_corrected"):
                                        auto_fixes.append(_r.get("function_name", "fn"))
                                except Exception:
                                    pass
                            fix_note = f" — 🔄 Code Tester auto-fixed: {', '.join(auto_fixes)}" if auto_fixes else ""
                            label = f"🔧 Executing tool: `{tool_str}` (call #{tool_call_count}){fix_note}…"
                            icon  = "🔧"

                        # Special-case: knowledge_agent — show retry count
                        if node_name == "knowledge_agent":
                            retry = node_output.get("kb_retry_count", 0)
                            if retry and retry > 1:
                                label = f"🔍 Rebuilding knowledge base (attempt {retry}/3)…"

                        # Special-case: critique_agent — show score if available
                        if node_name == "critique_agent":
                            critique = node_output.get("kb_critique", {})
                            if critique:
                                scores = critique.get("scores", {})
                                avg = round(sum(scores.values()) / 4, 1) if scores else "?"
                                status_word = "✅ Approved" if critique.get("status") == "approved" else "❌ Rejected"
                                label = f"🧐 Knowledge base critique complete — {status_word} · avg score {avg}/10"

                        step_placeholder.markdown(f"{icon} **{label}**")
                        run_status.update(label=f"{icon} {label}")

                        # Capture the final state from the last non-empty chunk
                        if node_output:
                            if final_state is None:
                                final_state = dict(initial_state)
                            final_state.update(node_output)

                run_status.update(label="✅ Analysis complete!", state="complete", expanded=False)
                step_placeholder.empty()

            except Exception as e:
                error_msg = str(e)
                run_status.update(label="❌ Analysis failed", state="error", expanded=True)
                st.error(f"An error occurred during analysis: {e}")

        if final_state and not error_msg:
            st.session_state["report"]       = final_state.get("report", "No report generated.")
            st.session_state["annotated_df"] = df
            st.session_state["agent_state"]  = final_state
            try:
                st.session_state["kg_html"] = kg_builder.build_knowledge_graph(df, "data/knowledge_graph.json")
            except Exception as kg_err:
                st.error(f"Error building knowledge graph: {kg_err}")
            st.session_state["analysis_done"] = True
            st.rerun()

    # If analysis is complete, show the report
    if st.session_state.get("analysis_done", False) and st.session_state["agent_state"]:
        state = st.session_state["agent_state"]
        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        st.markdown(ui_components.get_header_html("Analysis", "Report", "", "step 02 — quality report", ""), unsafe_allow_html=True)
        
        # Get actual issues count
        st.markdown("<div class='spacer-20'></div>", unsafe_allow_html=True)
        
        st.markdown('<div class="card"><div class="card-label">Executive AI Report</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="analysis-box">{st.session_state["report"]}</div></div>', unsafe_allow_html=True)
        
        st.markdown("<div class='spacer-20'></div>", unsafe_allow_html=True)
        
        col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])
        with col_dl1:
            try:
                docx_bytes = convert_report_to_docx(st.session_state["report"])
                st.download_button(
                    label="⬇ Download Report (.docx)",
                    data=docx_bytes,
                    file_name="data_quality_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as _docx_err:
                st.warning(f"Could not generate .docx: {_docx_err}")
        with col_dl2:
            st.download_button(
                label="⬇ Export Dataset (CSV)",
                data=convert_df_to_csv(st.session_state["annotated_df"]),
                file_name="analyzed_data.csv",
                mime="text/csv",
                use_container_width=True,
            )
        with col_dl3:
            try:
                with open("execute.py", "r", encoding="utf-8") as _ef:
                    _execute_src = _ef.read()
                st.download_button(
                    label="⬇ Download execute.py",
                    data=_execute_src.encode("utf-8"),
                    file_name="execute.py",
                    mime="text/x-python",
                    use_container_width=True,
                    help="Runnable Python script with all functions executed this session",
                )
            except Exception:
                st.info("execute.py will appear here once functions have been run.")
        
        st.markdown("<hr class='divider'>", unsafe_allow_html=True)
        st.markdown(ui_components.get_header_html("Follow-Up", "Chat Interface", "", "step 03 — conversational query", ""), unsafe_allow_html=True)
        st.markdown("<div class='spacer-20'></div>", unsafe_allow_html=True)
        
        # Render the chat dialog — each assistant turn may have an execution log expander
        for turn_idx, msg in enumerate(st.session_state.get("chat_history", [])):
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
            # After each assistant message, show execution log if one exists for that turn
            if msg["role"] == "assistant":
                exec_log = st.session_state["chat_execution_logs"].get(turn_idx)
                if exec_log:
                    with st.expander("📋 Execution Log — what the agent planned & ran", expanded=False):
                        # ── Plan: structured point-by-point display ───────────
                        if exec_log.get("plan"):
                            raw_plan = exec_log["plan"]

                            # Extract MODE badge
                            import re as _re
                            mode_match = _re.search(
                                r"MODE\s*[:\-]\s*(KB_ONLY|DATA_ONLY|BOTH|CONVERSATIONAL)",
                                raw_plan, _re.IGNORECASE
                            )
                            mode_val = mode_match.group(1).upper() if mode_match else "BOTH"
                            mode_colours = {
                                "KB_ONLY":       ("🟣", "Knowledge Base Only"),
                                "DATA_ONLY":     ("🟠", "Data Computation Only"),
                                "BOTH":          ("🔵", "KB + Data Computation"),
                                "CONVERSATIONAL":("⚪", "Conversational"),
                            }
                            mode_icon, mode_label = mode_colours.get(mode_val, ("🔵", mode_val))
                            st.markdown(f"**Execution Mode:** {mode_icon} `{mode_val}` — {mode_label}")
                            st.divider()

                            # Helper: extract a section between two headings
                            def _extract_section(text: str, heading: str) -> str:
                                pattern = rf"###\s*{_re.escape(heading)}\s*\n(.*?)(?=\n###|\Z)"
                                m = _re.search(pattern, text, _re.DOTALL | _re.IGNORECASE)
                                return m.group(1).strip() if m else ""

                            # UNDERSTANDING
                            understanding = _extract_section(raw_plan, "UNDERSTANDING")
                            if understanding:
                                st.markdown("**🧠 Understanding**")
                                for line in understanding.splitlines():
                                    line = line.strip().lstrip("- ").strip()
                                    if line:
                                        st.markdown(f"- {line}")

                            # EXECUTION STEPS
                            steps = _extract_section(raw_plan, "EXECUTION STEPS")
                            if steps:
                                st.markdown("**🪜 Execution Steps**")
                                for line in steps.splitlines():
                                    stripped = line.strip()
                                    if not stripped:
                                        continue
                                    if stripped.startswith("→"):
                                        # Sub-detail line (function info)
                                        st.markdown(f"  {stripped}")
                                    elif stripped.startswith("-"):
                                        st.markdown(stripped)
                                    else:
                                        st.markdown(f"- {stripped}")

                            # FUNCTION GAPS
                            gaps = _extract_section(raw_plan, "FUNCTION GAPS")
                            if gaps and "none" not in gaps.lower():
                                st.markdown("**🔩 Function Gaps (to be generated)**")
                                for line in gaps.splitlines():
                                    stripped = line.strip()
                                    if stripped and stripped not in ("-", "- None"):
                                        st.markdown(f"  {stripped}" if stripped.startswith("→") else stripped)

                            # EXPECTED OUTPUT
                            expected = _extract_section(raw_plan, "EXPECTED OUTPUT")
                            if expected:
                                st.markdown("**📤 Expected Output**")
                                for line in expected.splitlines():
                                    line = line.strip().lstrip("- ").strip()
                                    if line:
                                        st.markdown(f"- {line}")

                        # ── RAG chunks retrieved ──────────────────────────────
                        if exec_log.get("rag_chunks"):
                            st.divider()
                            st.markdown("**🔎 Knowledge Base Sections Retrieved (RAG)**")
                            raw = exec_log["rag_chunks"]
                            chunk_blocks = [b.strip() for b in raw.split("--- Chunk ") if b.strip()]
                            for block in chunk_blocks:
                                first_nl = block.find("\n")
                                header = block[:first_nl].rstrip(" -") if first_nl != -1 else block
                                body   = block[first_nl:].strip() if first_nl != -1 else ""
                                with st.expander(f"📄 {header}", expanded=False):
                                    st.markdown(body)

                        # ── Functions called ──────────────────────────────────
                        existing = exec_log.get("existing_functions", [])
                        new_funcs = exec_log.get("new_functions", [])
                        if existing or new_funcs:
                            st.divider()
                            st.markdown("**🔧 Functions Executed**")
                            if existing:
                                st.markdown("*Existing DB functions called:*")
                                for fn in existing:
                                    params_str = f" — params: `{fn['params']}`" if fn.get('params') else ""
                                    st.markdown(f"- ✅ `{fn['name']}`{params_str}")
                            if new_funcs:
                                st.markdown("*New functions generated & saved:*")
                                for fn in new_funcs:
                                    attempts = fn.get("tester_attempts", 1)
                                    auto_fixed = fn.get("auto_corrected", False)
                                    if auto_fixed:
                                        tester_badge = f"🔄 auto-corrected in {attempts} attempt{'s' if attempts != 1 else ''}"
                                    else:
                                        tester_badge = f"✔ passed in {attempts} attempt{'s' if attempts != 1 else ''}"
                                    st.markdown(
                                        f"- 🆕 `{fn['name']}` *(Group 4, pending approval)* "
                                        f"— Code Tester: {tester_badge}"
                                    )

                
        # Handle new prompts
        if prompt := st.chat_input("Ask a follow-up question about this dataset..."):
            st.session_state["chat_history"].append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
                
            with st.status("🧠 Analysing follow-up…", expanded=True) as chat_status:
                chat_placeholder = st.empty()
                from langchain_core.messages import HumanMessage
                import json as _json
                app = build_data_quality_graph(checkpointer=st.session_state["checkpointer"])
                config = {"configurable": {"thread_id": st.session_state["thread_id"]}}

                follow_state  = None
                planner_plan  = ""
                rag_chunks_text = ""
                existing_fns  = []   # {"name": str, "params": dict}
                new_fns       = []   # {"name": str, "description": str}

                for chunk in app.stream({"messages": [HumanMessage(content=prompt)]}, config, stream_mode="updates"):
                    for node_name, node_output in chunk.items():
                        icon, label = _NODE_LABELS.get(node_name, ("🔄", f"Running `{node_name}`…"))

                        # ── Planner: capture plan text ────────────────────────
                        if node_name == "chat_planner":
                            planner_plan = node_output.get("chat_plan", "")

                        # ── RAG: show how many chunks retrieved ───────────────
                        if node_name == "rag_retrieval":
                            rag_chunks_text = node_output.get("rag_chunks", "")
                            n_chunks = rag_chunks_text.count("--- Chunk ")
                            if n_chunks:
                                label = f"🔎 Retrieved {n_chunks} relevant KB sections via RAG…"
                            else:
                                label = "🔎 RAG retrieval skipped (DATA_ONLY or CONVERSATIONAL mode)"

                        # ── Tool execution: track which functions ran ─────────
                        if node_name == "tool_execution":
                            msgs = node_output.get("messages", [])
                            tool_names_set = set()
                            for m in msgs:
                                raw_name = m.get("name", "") if isinstance(m, dict) else getattr(m, "name", "")
                                tool_names_set.add(raw_name)
                                # Parse the tool result JSON for function metadata
                                raw_content = m.get("content", "") if isinstance(m, dict) else getattr(m, "content", "")
                                try:
                                    res = _json.loads(raw_content) if isinstance(raw_content, str) else {}
                                except Exception:
                                    res = {}
                                if raw_name == "generate_and_test_custom_function":
                                    fn_name = res.get("function_name", "unknown")
                                    auto_fixed = res.get("code_auto_corrected", False)
                                    tester_attempts = res.get("code_tester_attempts", 1)
                                    new_fns.append({
                                        "name": fn_name,
                                        "description": res.get("database_status", ""),
                                        "auto_corrected": auto_fixed,
                                        "tester_attempts": tester_attempts,
                                    })
                                elif raw_name == "execute_existing_function_with_params":
                                    existing_fns.append({
                                        "name": res.get("function_name", "unknown"),
                                        "params": res.get("params_used", {})
                                    })
                            label = f"🔧 Executing tool: `{', '.join(tool_names_set) or 'custom function'}`…"
                            icon  = "🔧"

                        chat_placeholder.markdown(f"{icon} **{label}**")
                        chat_status.update(label=f"{icon} {label}")
                        if node_output:
                            follow_state = node_output

                chat_status.update(label="✅ Done", state="complete", expanded=False)
                chat_placeholder.empty()

                ai_response = (
                    follow_state["messages"][-1].content
                    if follow_state and follow_state.get("messages")
                    else "No response."
                )
                st.session_state["chat_history"].append({"role": "assistant", "content": ai_response})

                # ── Persist execution log keyed to this assistant turn ────────
                turn_idx = len(st.session_state["chat_history"]) - 1
                if planner_plan or existing_fns or new_fns or rag_chunks_text:
                    # Extract mode from planner output for badge display
                    import re as _re2
                    _mode_m = _re2.search(
                        r"MODE\s*[:\-]\s*(KB_ONLY|DATA_ONLY|BOTH|CONVERSATIONAL)",
                        planner_plan, _re2.IGNORECASE
                    )
                    stored_mode = _mode_m.group(1).upper() if _mode_m else "BOTH"
                    st.session_state["chat_execution_logs"][turn_idx] = {
                        "plan":               planner_plan,
                        "chat_mode":          stored_mode,
                        "rag_chunks":         rag_chunks_text,
                        "existing_functions": existing_fns,
                        "new_functions":      new_fns,
                    }
                st.rerun()

with tab2:
    st.header("Internal Functions Database")
    st.markdown("View all pre-defined and agent-generated data quality functions securely stored in `database/app.db`.")
    
    try:
        import sqlite3
        with sqlite3.connect("database/app.db", timeout=5.0) as conn:
            query = "SELECT function_name, function_description, function_group, approved_by_team, updated_at FROM data_quality_functions"
            db_df = pd.read_sql_query(query, conn)
            
            # Convert boolean column to string for better display
            if 'approved_by_team' in db_df.columns:
                db_df['approved_by_team'] = db_df['approved_by_team'].apply(lambda x: "Yes" if x == 1 else "No (Quarantined)")
            
            # Map function groups to human-readable labels
            if 'function_group' in db_df.columns:
                group_map = {
                    1: "Group 1 (Dataset-Level)", 
                    2: "Group 2 (Metadata-Dependent)", 
                    3: "Group 3 (Domain Logic)", 
                    4: "Group 4 (AI Generated)"
                }
                db_df['function_group'] = db_df['function_group'].map(group_map).fillna("Unknown")
            
            # Display stats
            st.metric("Total Functions inside DB", len(db_df))
            
            # Full table display
            st.dataframe(
                db_df,
                use_container_width=True,
                hide_index=True
            )
    except Exception as e:
        st.error(f"Could not load database: {e}")

with tab3:
    st.header("Dataset Knowledge Graph")
    st.markdown("This interactive graph visualizes the strongest mathematical correlations (top 5 per feature) found across all numeric columns in your dataset. Hover over nodes to see their exact Min and Max values.")
    
    if st.session_state.get("kg_html"):
        components.html(st.session_state["kg_html"], height=650)
        
        try:
            with open("data/knowledge_graph.json", "rb") as f:
                kg_json = f.read()
            st.download_button(
                label="⬇ Download Knowledge Graph (JSON)", 
                data=kg_json, 
                file_name="knowledge_graph.json", 
                mime="application/json"
            )
        except Exception:
            pass
    else:
        st.info("No Knowledge Graph available yet. Please upload a dataset and run the Quality Analysis on the New Analysis tab.")

with tab4:
    st.header("Domain Knowledge Base")
    st.markdown("View all the Process, Physics, Equipment, and OEM knowledge constraints automatically injected into the Reasoning Agent prior to analysis.")
    
    try:
        import sqlite3
        with sqlite3.connect("database/app.db", timeout=5.0) as conn:
            current_thread = st.session_state.get("thread_id", "")
            query = "SELECT category, topic, knowledge_text, updated_at FROM domain_knowledge WHERE thread_id = ? ORDER BY category, topic"
            kb_df = pd.read_sql_query(query, conn, params=(current_thread,))
            
            if len(kb_df) == 0:
                query_all = "SELECT category, topic, knowledge_text, updated_at, thread_id FROM domain_knowledge ORDER BY category, topic"
                kb_df = pd.read_sql_query(query_all, conn)
                st.info("Displaying the **Global Knowledge Base** aggregator. Start a new analysis to drill down into a thread-specific constraints block.")
            
            st.metric("Total Knowledge Entries", len(kb_df))
            st.dataframe(
                kb_df,
                use_container_width=True,
                hide_index=True
            )
    except Exception as e:
        st.error(f"Could not load knowledge database: {e}")

with tab5:
    st.header("Agent Workflow Diagram")
    st.markdown("Interactive graph of the LangGraph pipeline. **Drag nodes**, scroll to zoom, hover for descriptions.")
    components.html(ui_components.get_workflow_graph_html(), height=680)
